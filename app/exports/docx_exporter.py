from __future__ import annotations

import os
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app import config
from app.services.assessment_service import AssessmentView
from app.services.settings_service import AppSettings
from app.utils.helpers import format_datetime, report_basename

_PRIMARY = RGBColor(0x10, 0x1A, 0x33)
_ACCENT = RGBColor(0x2D, 0x6C, 0xDF)
_TEXT_DARK = RGBColor(0x1B, 0x1F, 0x2A)
_TEXT_MUTED = RGBColor(0x55, 0x5C, 0x6B)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_HEADER_FILL = "13213F"
_ZEBRA_FILL = "F1F4FB"
_LABEL_FILL = "E5EAF6"

_SEVERITY_FILL = {
    "Critical": "C0182B",
    "High": "E0561B",
    "Medium": "C9911C",
    "Low": "2E8B57",
    "Informational": "2D6CDF",
}

_FONT = "Segoe UI"


def _set_cell_background(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, left: int = 120, right: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for tag, value in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def _disable_autofit(table) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _style_run(run, *, size: float = 10.5, bold: bool = False, color: RGBColor = _TEXT_DARK) -> None:
    run.font.name = _FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _FONT)
    rfonts.set(qn("w:hAnsi"), _FONT)


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10.5,
    color: RGBColor = _TEXT_DARK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    lines = text.split("\n") if text else [""]
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        _style_run(run, size=size, bold=bold, color=color)
        if index < len(lines) - 1:
            run.add_break()


def _section_heading(document, title: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _disable_autofit(table)
    table.columns[0].width = Cm(17.0)
    cell = table.cell(0, 0)
    cell.width = Cm(17.0)
    _set_cell_background(cell, _HEADER_FILL)
    _set_cell_margins(cell, top=90, bottom=90, left=160, right=160)
    _write_cell(cell, title.upper(), bold=True, size=11.5, color=_WHITE)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _key_value_table(document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _disable_autofit(table)
    label_width = Cm(5.4)
    value_width = Cm(11.6)
    for index, (label, value) in enumerate(rows):
        label_cell = table.cell(index, 0)
        value_cell = table.cell(index, 1)
        label_cell.width = label_width
        value_cell.width = value_width
        _set_cell_background(label_cell, _LABEL_FILL)
        if index % 2 == 1:
            _set_cell_background(value_cell, _ZEBRA_FILL)
        _set_cell_margins(label_cell)
        _set_cell_margins(value_cell)
        _write_cell(label_cell, label, bold=True, size=10, color=_PRIMARY)
        _write_cell(value_cell, value if value else "-", size=10, color=_TEXT_DARK)
    _set_table_borders(table)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _matrix_table(document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _disable_autofit(table)
    for column_index, header in enumerate(headers):
        cell = table.cell(0, column_index)
        cell.width = Cm(widths[column_index])
        _set_cell_background(cell, _HEADER_FILL)
        _set_cell_margins(cell)
        _write_cell(cell, header, bold=True, size=10, color=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_index, row in enumerate(rows, start=1):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.width = Cm(widths[column_index])
            if row_index % 2 == 0:
                _set_cell_background(cell, _ZEBRA_FILL)
            _set_cell_margins(cell)
            align = WD_ALIGN_PARAGRAPH.CENTER if column_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            _write_cell(cell, value, size=10, color=_TEXT_DARK, align=align)
    _set_table_borders(table)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _bullet_table(document, title: str, items: list[str]) -> None:
    table = document.add_table(rows=len(items) + 1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _disable_autofit(table)
    header = table.cell(0, 0)
    header.width = Cm(17.0)
    _set_cell_background(header, _ACCENT_HEX_FILL)
    _set_cell_margins(header)
    _write_cell(header, title, bold=True, size=10, color=_WHITE)
    for index, item in enumerate(items, start=1):
        cell = table.cell(index, 0)
        cell.width = Cm(17.0)
        if index % 2 == 0:
            _set_cell_background(cell, _ZEBRA_FILL)
        _set_cell_margins(cell)
        _write_cell(cell, f"{index}.  {item}", size=10, color=_TEXT_DARK)
    _set_table_borders(table)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


_ACCENT_HEX_FILL = "2D6CDF"


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "C7D0E4")
        borders.append(element)
    tbl_pr.append(borders)


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    _style_run(run, size=8.5, color=_TEXT_MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(fld_end)


def _configure_section(section, settings: AppSettings) -> None:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)

    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    _clear_paragraph(header_paragraph)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_text = settings.report_header or "VULNERABILITY ASSESSMENT REPORT"
    run = header_paragraph.add_run(header_text)
    _style_run(run, size=8.5, bold=True, color=_TEXT_MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    _clear_paragraph(footer_paragraph)
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_label = settings.report_footer or f"{settings.institution_name} | {config.APP_NAME} v{config.APP_VERSION}"
    left_run = footer_paragraph.add_run(f"{footer_label}    |    Halaman ")
    _style_run(left_run, size=8.5, color=_TEXT_MUTED)
    _add_page_number_field(footer_paragraph)


def _build_cover(document, view: AssessmentView, settings: AppSettings) -> None:
    for _ in range(2):
        document.add_paragraph()

    if settings.report_logo_path and os.path.isfile(settings.report_logo_path):
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            logo_paragraph.add_run().add_picture(settings.report_logo_path, width=Cm(3.6))
        except (OSError, ValueError):
            pass

    institution = document.add_paragraph()
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    institution.paragraph_format.space_before = Pt(10)
    institution_run = institution.add_run(settings.institution_name or config.APP_NAME)
    _style_run(institution_run, size=15, bold=True, color=_PRIMARY)

    if settings.institution_address:
        address = document.add_paragraph()
        address.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_run = address.add_run(settings.institution_address)
        _style_run(address_run, size=10, color=_TEXT_MUTED)

    document.add_paragraph()
    document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(settings.report_header or "VULNERABILITY ASSESSMENT REPORT")
    _style_run(title_run, size=24, bold=True, color=_ACCENT)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(4)
    subtitle_run = subtitle.add_run("Laporan Penilaian Akhir Kerentanan")
    _style_run(subtitle_run, size=12, color=_TEXT_MUTED)

    document.add_paragraph()

    badge_table = document.add_table(rows=1, cols=1)
    badge_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _disable_autofit(badge_table)
    badge_cell = badge_table.cell(0, 0)
    badge_cell.width = Cm(9.0)
    fill = _SEVERITY_FILL.get(view.final_severity, "2D6CDF")
    _set_cell_background(badge_cell, fill)
    _set_cell_margins(badge_cell, top=120, bottom=120, left=160, right=160)
    _write_cell(
        badge_cell,
        f"{view.final_severity.upper()}  |  CVSS {view.cvss_score:.1f}",
        bold=True,
        size=14,
        color=_WHITE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    for _ in range(4):
        document.add_paragraph()

    meta_rows = [
        ("Target / URL", view.target),
        ("Nama Temuan", view.finding_name),
        ("Tanggal Ditemukan", view.found_on.strftime("%d %B %Y")),
        ("Peneliti / Handle", view.researcher),
        ("Tanggal Laporan", datetime.now().strftime("%d %B %Y")),
    ]
    _key_value_table(document, meta_rows)


def _build_body(document, view: AssessmentView, settings: AppSettings) -> None:
    _section_heading(document, "Informasi Temuan")
    _key_value_table(
        document,
        [
            ("Target / URL", view.target),
            ("Nama Temuan", view.finding_name),
            ("Tanggal Ditemukan", view.found_on.strftime("%d %B %Y")),
            ("Peneliti / Handle", view.researcher),
            ("Pemilik Penilaian", view.owner_username),
        ],
    )

    _section_heading(document, "CVSS Scoring")
    _key_value_table(
        document,
        [
            ("Vector String", view.cvss_vector),
            ("Base Score", f"{view.cvss_score:.1f}"),
            ("Severity Awal", view.base_severity),
            ("Severity Akhir", view.final_severity),
            ("Severity Upgrade", "Ya" if view.severity_upgraded else "Tidak"),
        ],
    )

    _section_heading(document, "Parameter Detail")
    parameter_rows = [
        [code, view.parameter_labels.get(label, "-"), label]
        for label, code in (
            ("Attack Vector", view.metric_codes["AV"]),
            ("Attack Complexity", view.metric_codes["AC"]),
            ("Privileges Required", view.metric_codes["PR"]),
            ("User Interaction", view.metric_codes["UI"]),
            ("Scope", view.metric_codes["S"]),
            ("Confidentiality Impact", view.metric_codes["C"]),
            ("Integrity Impact", view.metric_codes["I"]),
            ("Availability Impact", view.metric_codes["A"]),
        )
    ]
    _matrix_table(
        document,
        ["Kode", "Nilai", "Parameter"],
        [[row[0], row[1], row[2]] for row in parameter_rows],
        [2.6, 6.0, 8.4],
    )

    _section_heading(document, "Exploitability")
    _key_value_table(
        document,
        [
            ("Public Exploit / CVE", "Ya" if view.public_exploit else "Tidak"),
            ("CVE ID", view.cve_id or "-"),
            ("Reproducible", "Ya" if view.reproducible else "Tidak"),
            ("Proof of Concept", "Tersedia" if view.poc_available else "Tidak tersedia"),
            ("Lampiran PoC", os.path.basename(view.poc_attachment_path) if view.poc_attachment_path else "-"),
        ],
    )

    _section_heading(document, "Impact Detail")
    _key_value_table(
        document,
        [
            ("Data Terdampak", view.impacted_data),
            ("Estimasi Jumlah User", view.estimated_users),
            ("Regulasi Terkait", ", ".join(view.regulations) if view.regulations else "-"),
            ("Business Impact", view.business_impact),
        ],
    )

    _section_heading(document, "Rekomendasi Remediation")
    _bullet_table(document, f"Langkah Teknis - {view.remediation_name}", list(view.remediation_steps))
    _key_value_table(
        document,
        [
            ("Referensi OWASP", view.remediation_owasp),
            ("CWE", view.remediation_cwe),
            ("CAPEC", view.remediation_capec),
            ("Target Deadline", view.deadline_label),
        ],
    )
    if view.remediation_references:
        _bullet_table(document, "Referensi Tambahan", list(view.remediation_references))

    _section_heading(document, "Catatan Severity Chaining")
    if view.severity_upgraded and view.severity_reasons:
        _bullet_table(document, "Alasan Peningkatan Severity", list(view.severity_reasons))
    else:
        note = document.add_paragraph()
        note.paragraph_format.space_before = Pt(2)
        note_run = note.add_run(
            "Tidak ditemukan kombinasi kondisi yang memicu peningkatan severity. "
            "Severity akhir setara dengan hasil perhitungan CVSS dasar."
        )
        _style_run(note_run, size=10, color=_TEXT_MUTED)


def export_docx(view: AssessmentView, settings: AppSettings, output_dir: str | None = None) -> str:
    config.ensure_directories()
    target_dir = output_dir or str(config.EXPORT_DIR)
    os.makedirs(target_dir, exist_ok=True)

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10.5)

    section = document.sections[0]
    _configure_section(section, settings)
    _build_cover(document, view, settings)

    document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(document.sections[-1], settings)
    _build_body(document, view, settings)

    basename = report_basename(view.target, view.finding_name, view.found_on)
    output_path = os.path.join(target_dir, f"{basename}.docx")
    document.save(output_path)
    return output_path
